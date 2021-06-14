import docx

doc = docx.Document('document-1.docx')
print(doc.core_properties.title)
print(doc.core_properties.keywords)
print(doc.core_properties.modified)
print(len(doc.paragraphs))

for index, paragraph in enumerate(doc.paragraphs):
    if paragraph.text:
        print(index, paragraph.text)

